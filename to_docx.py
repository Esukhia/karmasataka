from docx import Document  # package name: python-docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Length, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from horology import timed


#@timed(unit="min")
def create_docx(chunks, path):
    # CONFIG
    indentation = 0.5
    tibetan_style = {
        'font': 'Jomolhari',
        'size': 8
    }
    pedurma_style = {
        'color': (112, 128, 144),
        'font': 'Jomolhari',
        'size': 5
    }
    semantic_style = {
        'font': 'Free Mono',
        'size': 8
    }
    communicative_style = {
        'font': 'Liberation Serif',
        'size': 12
    }

    document = Document()
    styles = document.styles

    # TIBETAN
    bo_style = styles.add_style('Tibetan', WD_STYLE_TYPE.CHARACTER)
    # bo_style.base_style = styles['Normal']
    bo_font = bo_style.font
    bo_font.name = tibetan_style['font']
    bo_font.size = Pt(tibetan_style['size'])
    # PEYDURMA NOTES
    note_style = styles.add_style('Peydurma Notes', WD_STYLE_TYPE.CHARACTER)
    # note_style.base_style = styles['Normal']
    note_font = note_style.font
    note_font.name = pedurma_style['font']
    note_font.size = Pt(pedurma_style['size'])
    # note_font.subscript = True
    c = pedurma_style['color']
    note_font.color.rgb = RGBColor(c[0], c[1], c[2])

    # COMMUNICATIVE VERSION
    com_style = styles.add_style('Communicative', WD_STYLE_TYPE.CHARACTER)
    # com_style.base_style = styles['Normal']
    com_font = com_style.font
    com_font.name = communicative_style['font']
    com_font.size = Pt(communicative_style['size'])

    # SEMANTIC VERSION
    sem_style = styles.add_style('Semantic', WD_STYLE_TYPE.CHARACTER)
    sem_style.base_style = styles['Normal']
    sem_font = sem_style.font
    sem_font.name = semantic_style['font']
    sem_font.size = Pt(semantic_style['size'])

    for chunk in chunks:
        com, others = chunk
        p = document.add_paragraph()
        #########################
        # Communicative
        # p = document.add_paragraph()
        # par_format = p.paragraph_format
        # par_format.space_before = Pt(0)
        run = p.add_run(com)
        run.style = 'Communicative'
        run.add_break()
        # ************************

        for pair in others:
            bo, sem = pair
            bo = re.split('(<.+?>)', bo)
            for b in bo:
                run = p.add_run(b)
                if b.startswith('<'):
                    run.style = 'Peydurma Notes'
                else:
                    run.style = 'Tibetan'
            p.add_run().add_break()

            run = p.add_run(sem)
            run.style = 'Semantic'
            run.add_break()

        # par_format = p.paragraph_format
        # par_format.left_indent = Cm(indentation)
        # par_format.space_after = Pt(0)
        #

    out_path = path.parent / (path.stem + '.docx')
    document.save(str(out_path))
